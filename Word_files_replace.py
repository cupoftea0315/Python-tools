"""
        本项目特点：
            1、获取word中数据,批量替换word中文字！

"""
import os
from docx import Document


def replace_word(doc, old_word, new_word):
    """
    定义批量替换文字的函数
    :param doc: 要替换的文档
    :param old: 被替换的文字
    :param new: 替换后的文字
    :return:
    """
    # ss = False
    for p in doc.paragraphs:  # 遍历文档段落

        for run in p.runs:  # 遍历段落的字块
            # if "一般公共预算当年拨款结构情况" in run.text:
            # ss = True
            # elif "一般公共预算当年拨款具体使用情况" in run.text:
            # ss = False
            # if ss:
            # if "（类）" not in run.text:
            if old_word in run.text:
                run.text = run.text.replace(old_word, new_word)

    # 遍历文档的表格， 替换表格里的要替换的文字
    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             cell.text = cell.text.replace(old_word, new_word)


data_text = [
    {"old": "自然地理", "new": "physical geography"},
]
for i in range(100):
    try:
        path = input(str(i + 1) + ':请输入文件夹地址：').replace('\\', '/')

        for file_name in os.listdir(path):

            print(file_name)
            doc_path = path + '/' + file_name
            doc = Document(doc_path)
            for data in data_text:
                replace_word(doc, data["old"], data["new"])
            doc.save(file_name)
        print("================== 完成个文件 =================")
        input('=============================================')
    except Exception as E:
        print('错误：', '1.文件夹里的文件不是”.docx“类型文件！  2.文件地址有无（”D:\”）')
